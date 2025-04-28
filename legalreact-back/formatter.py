import os
import json
import re
import logging
from azure.storage.blob import BlobServiceClient
from langchain_openai import AzureChatOpenAI
from docx import Document
from difflib import get_close_matches
from llm import LLM

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Formatter:
    def __init__(self):
        self.AZURE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        self.CONTAINER_NAME = os.getenv("AZURE_CONTAINER_NAME_4")
        self.llm = LLM()
        self.gen_llm = self.llm.initialize_gen_llm()


    def list_templates_from_blob(self):
        try:
            blob_service_client = BlobServiceClient.from_connection_string(self.AZURE_CONNECTION_STRING)
            container_client = blob_service_client.get_container_client(self.CONTAINER_NAME)
            templates = [os.path.splitext(os.path.basename(blob.name))[0]
                        for blob in container_client.list_blobs() if blob.name.endswith(".docx")]
            return templates
        except Exception as e:
            logging.error(f"Error listing templates: {e}")
            return []


    def classify_document_type(self, user_query):
        try:
            response = self.gen_llm.invoke([
                {"role": "system", "content": "Identify the most appropriate document type that is being asked for in the user query. Only return the document type name with no extra text. If it is a non disclosure agreement then return NDA. Valid template names should have NDA or Business Partnership in them."},
                {"role": "user", "content": user_query}
            ])
            if response and hasattr(response, 'content'):
                doc_type = response.content.strip().upper().replace(" ", "_")
                logging.info(f"Classified Document Type: {doc_type}")
                return doc_type
            logging.error("Document classification failed.")
            return None
        except Exception as e:
            logging.error(f"Error in document classification: {e}")
            return None


    def fetch_template_from_blob(self, document_type):
        try:
            blob_service_client = BlobServiceClient.from_connection_string(self.AZURE_CONNECTION_STRING)
            container_client = blob_service_client.get_container_client(self.CONTAINER_NAME)
            templates = [blob.name.replace(".docx", "") for blob in container_client.list_blobs()]
            best_match = get_close_matches(document_type, templates, n=1, cutoff=0.5)
            if not best_match:
                logging.error(f"No matching template found for {document_type}")
                return None

            selected = best_match[0]
            blob_client = blob_service_client.get_blob_client(self.CONTAINER_NAME, f"{selected}.docx")
            os.makedirs("templates", exist_ok=True)
            local_path = f"templates/{selected}.docx"

            with open(local_path, "wb") as file:
                file.write(blob_client.download_blob().readall())

            logging.info(f"Template downloaded: {selected}.docx")
            return local_path
        except Exception as e:
            logging.error(f"Error fetching template: {e}")
            return None


    def extract_placeholders(self, template_path):
        try:
            doc = Document(template_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return list(set(re.findall(r"\{(.*?)\}", text)))
        except Exception as e:
            logging.error(f"Error extracting placeholders: {e}")
            return []


    def generate_extraction_prompt(self, user_query, document_type, placeholders):
        return f"""
        The input is: {user_query}. The following is a description of a {document_type}. 

        Extract and classify:
        1. **DISCLOSING_PARTIES**
        2. **RECEIVING_PARTIES**
        3. **AGREEMENT_DATE** (YYYY-MM-DD)
        4. **COMMENCEMENT_DATE** (YYYY-MM-DD)
        5. **TERM_YEARS**

        Format the output strictly in JSON:
        {{
            "AGREEMENT_DATE": "YYYY-MM-DD",
            "COMMENCEMENT_DATE": "YYYY-MM-DD",
            "TERM_YEARS": "X",
            "DISCLOSING_PARTIES": ["Name1"],
            "RECEIVING_PARTIES": ["Name2"]
        }}
        """


    def extract_entities_from_gpt(self, user_query, document_type, placeholders):
        prompt = self.generate_extraction_prompt(user_query, document_type, placeholders)
        response = self.gen_llm.invoke([{"role": "user", "content": prompt}])
        return self.extract_json_from_response(response.content.strip())


    def extract_json_from_response(self, response_text):
        try:
            response_text = re.sub(r"```json\n|\n```", "", response_text).strip()
            logging.info(f"GPT Response: {response_text}")
            return json.loads(response_text)
        except json.JSONDecodeError:
            logging.warning("JSON decoding failed.")
            return None


    def fill_document_with_gpt(self, template_path, extracted_data):
        try:
            doc = Document(template_path)
            cleaned = {k: (", ".join(v) if isinstance(v, list) else str(v or "")) for k, v in extracted_data.items()}

            for para in doc.paragraphs:
                for placeholder, value in cleaned.items():
                    para.text = para.text.replace(f"{{ {placeholder} }}", value)

            output_path = "/Users/aryan_zingade/Downloads/generated_document.docx"
            doc.save(output_path)
            logging.info(f"Document saved: {output_path}")
            return output_path
        except Exception as e:
            logging.error(f"Error filling document: {e}")
            return None


    def process_document(self, user_query):
        logging.info("Listing templates from blob...")
        if not self.list_templates_from_blob():
            return None

        logging.info("Classifying document type...")
        doc_type = self.classify_document_type(user_query)
        if not doc_type:
            return None

        logging.info("Fetching template from blob...")
        template_path = self.fetch_template_from_blob(doc_type)
        if not template_path:
            return None

        logging.info("Extracting placeholders...")
        placeholders = self.extract_placeholders(template_path)
        if not placeholders:
            return None

        logging.info("Extracting entities using GPT...")
        data = self.extract_entities_from_gpt(user_query, doc_type, placeholders)
        if not data:
            return None

        logging.info("Filling the document with extracted data...")
        return self.fill_document_with_gpt(template_path, data)
